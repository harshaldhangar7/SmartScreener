#!/usr/bin/env python3
"""
Bulk Upload Script for SmartScreener
Uploads all resume and job description files from the project directories
"""

import os
import sys
import glob
import time
from pathlib import Path

# Add the current directory to Python path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def bulk_upload_resumes():
    """Upload all resume files from the resumes directory"""
    from app import app, db, Candidate
    from resume_parser import parse_resume
    from ai_resume_analyzer import ai_analyzer
    from sentence_transformers import SentenceTransformer
    import base64
    import json
    import time
    from datetime import datetime

    resume_dir = Path("resumes")
    if not resume_dir.exists():
        print("[ERROR] Resumes directory not found!")
        return 0

    resume_files = list(resume_dir.glob("*"))
    resume_files = [f for f in resume_files if f.suffix.lower() in ['.pdf', '.docx', '.doc']]

    if not resume_files:
        print("[ERROR] No resume files found in resumes/ directory")
        return 0

    print(f"[INFO] Found {len(resume_files)} resume files to upload")

    uploaded_count = 0
    skipped_count = 0

    # Load embedding model
    try:
        embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    except Exception as e:
        print(f"[ERROR] Failed to load embedding model: {e}")
        return 0

    with app.app_context():
        for resume_file in resume_files:
            try:
                print(f"Uploading: {resume_file.name}")

                # Parse the resume
                candidate_data = parse_resume(str(resume_file), resume_file.name)

                # Validate parsed data
                if not candidate_data.get('name'):
                    print(f"[WARNING] Could not extract candidate name from {resume_file.name}")
                    skipped_count += 1
                    continue

                # Create new candidate
                new_candidate = Candidate(
                    name=candidate_data.get('name', ''),
                    email=candidate_data.get('email', ''),
                    phone=candidate_data.get('phone', ''),
                    skills=candidate_data.get('skills', []),
                    experience_years=float(candidate_data.get('experience_years', 0)),
                    education=candidate_data.get('education', []),
                    experience_entries=candidate_data.get('experience_entries', []),
                    resume_filename=resume_file.name,
                    parsed_date=datetime.now()
                )

                # Save candidate data to database
                db.session.add(new_candidate)
                db.session.commit()

                # Perform AI analysis if available
                if ai_analyzer and ai_analyzer.is_available():
                    try:
                        # Extract raw text for AI analysis
                        raw_text = ""
                        try:
                            if resume_file.suffix.lower() == '.pdf':
                                import fitz
                                doc = fitz.open(str(resume_file))
                                for page in doc:
                                    raw_text += page.get_text()
                            elif resume_file.suffix.lower() in ['.docx', '.doc']:
                                import docx
                                doc = docx.Document(str(resume_file))
                                for paragraph in doc.paragraphs:
                                    raw_text += paragraph.text + '\n'
                        except Exception as e:
                            print(f"[WARNING] Failed to extract raw text for AI analysis: {e}")

                        ai_analysis = ai_analyzer.analyze_resume_comprehensive(raw_text, candidate_data)
                        ai_summary = ai_analyzer.generate_candidate_summary(candidate_data, ai_analysis)

                        # Update candidate with AI analysis
                        new_candidate.ai_analysis = ai_analysis
                        new_candidate.ai_summary = ai_summary
                        new_candidate.ai_quality_score = ai_analysis.get('overall_quality_score', 0)
                        new_candidate.ai_analysis_date = datetime.now()

                        db.session.commit()
                    except Exception as e:
                        print(f"[WARNING] AI analysis failed for {resume_file.name}: {e}")

                # Compute and store candidate embedding
                try:
                    text_parts = [
                        new_candidate.name or '',
                        ', '.join(new_candidate.skills or []),
                        ', '.join([exp.get('title','') for exp in (new_candidate.experience_entries or []) if isinstance(exp, dict)]),
                        ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (new_candidate.education or [])])
                    ]
                    text_value = ' \n '.join([p for p in text_parts if p])
                    vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()

                    def encode_vector_to_text(vector):
                        try:
                            return base64.b64encode(json.dumps(vector).encode("utf-8")).decode("utf-8")
                        except Exception:
                            return json.dumps(vector)

                    new_candidate.embedding = encode_vector_to_text(vector)
                    db.session.commit()
                except Exception as e:
                    print(f"[WARNING] Failed to compute embedding for {resume_file.name}: {e}")

                print(f"[SUCCESS] Successfully uploaded: {resume_file.name}")
                uploaded_count += 1

                # Small delay to avoid overwhelming the system
                time.sleep(0.1)

            except Exception as e:
                db.session.rollback()
                print(f"[ERROR] Error uploading {resume_file.name}: {str(e)}")
                skipped_count += 1

                # Small delay to avoid overwhelming the system
                time.sleep(0.5)

            except Exception as e:
                print(f"❌ Error processing {resume_file.name}: {str(e)}")
                skipped_count += 1

    return uploaded_count, skipped_count

def bulk_upload_job_descriptions():
    """Upload all job description files from the job_descriptions directory"""
    from app import app, db, JobDescription, parse_job_description
    from sentence_transformers import SentenceTransformer
    import base64
    import json
    import time
    from datetime import datetime

    job_dir = Path("job_descriptions")
    if not job_dir.exists():
        print("[ERROR] Job descriptions directory not found!")
        return 0

    job_files = list(job_dir.glob("*"))
    job_files = [f for f in job_files if f.suffix.lower() in ['.pdf', '.docx', '.doc']]

    if not job_files:
        print("[ERROR] No job description files found in job_descriptions/ directory")
        return 0

    print(f"[INFO] Found {len(job_files)} job description files to upload")

    uploaded_count = 0
    skipped_count = 0

    # Load embedding model
    try:
        embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    except Exception as e:
        print(f"[ERROR] Failed to load embedding model: {e}")
        return 0

    with app.app_context():
        for job_file in job_files:
            try:
                print(f"Uploading: {job_file.name}")

                # Extract text from the job description file
                job_text = ""
                try:
                    if job_file.suffix.lower() == '.pdf':
                        import fitz
                        doc = fitz.open(str(job_file))
                        for page in doc:
                            job_text += page.get_text()
                    elif job_file.suffix.lower() in ['.docx', '.doc']:
                        import docx
                        doc = docx.Document(str(job_file))
                        for paragraph in doc.paragraphs:
                            job_text += paragraph.text + '\n'
                except Exception as e:
                    print(f"[WARNING] Failed to extract text from {job_file.name}: {e}")
                    skipped_count += 1
                    continue

                # Parse the job description
                job_data = parse_job_description(job_text, job_file.name)

                # Create new job description
                new_job = JobDescription(
                    title=job_data.get('title', f'Job Description from {job_file.name}'),
                    required_skills=job_data.get('required_skills', []),
                    min_experience=float(job_data.get('min_experience', 0)),
                    required_education=job_data.get('required_education', []),
                    fk_user_id=1  # Assuming user ID 1, you may need to adjust this
                )

                # Save job description to database
                db.session.add(new_job)
                db.session.commit()

                # Compute and store job embedding
                try:
                    text_parts = [
                        new_job.title or '',
                        ', '.join(new_job.required_skills or []),
                        ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (new_job.required_education or [])])
                    ]
                    text_value = ' \n '.join([p for p in text_parts if p])
                    vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()

                    def encode_vector_to_text(vector):
                        try:
                            return base64.b64encode(json.dumps(vector).encode("utf-8")).decode("utf-8")
                        except Exception:
                            return json.dumps(vector)

                    new_job.embedding = encode_vector_to_text(vector)
                    db.session.commit()
                except Exception as e:
                    print(f"[WARNING] Failed to compute embedding for {job_file.name}: {e}")

                print(f"[SUCCESS] Successfully uploaded: {job_file.name}")
                uploaded_count += 1

                # Small delay to avoid overwhelming the system
                time.sleep(0.1)

            except Exception as e:
                db.session.rollback()
                print(f"[ERROR] Error uploading {job_file.name}: {str(e)}")
                skipped_count += 1

                # Small delay to avoid overwhelming the system
                time.sleep(0.5)

            except Exception as e:
                print(f"❌ Error processing {job_file.name}: {str(e)}")
                skipped_count += 1

    return uploaded_count, skipped_count

def main():
    """Main function to run bulk upload"""
    import argparse

    parser = argparse.ArgumentParser(description="Bulk upload resumes and job descriptions")
    parser.add_argument("--skip-ai", action="store_true", help="Skip AI analysis to avoid rate limits")
    args = parser.parse_args()

    print("=== SmartScreener Bulk Upload Script ===")
    print("=" * 50)

    if args.skip_ai:
        print("[INFO] AI analysis will be skipped to avoid rate limits")
        # Temporarily disable AI analyzer
        global ai_analyzer
        ai_analyzer._client = None
    else:
        print("[INFO] AI analysis is enabled (may hit rate limits with many files)")

    # Check if we're in the right directory
    if not Path("app.py").exists():
        print("[ERROR] Please run this script from the SmartScreener project root directory")
        sys.exit(1)

    # Check if database exists
    db_path = Path("instance/smartscreener.db")
    if not db_path.exists():
        print("[ERROR] Database not found. Please run the application first to create the database.")
        sys.exit(1)

    print("[INFO] Starting bulk upload process...\n")

    # Upload resumes
    print("UPLOADING RESUMES")
    print("-" * 30)
    resume_uploaded, resume_skipped = bulk_upload_resumes()
    print(f"[SUCCESS] Resumes uploaded: {resume_uploaded}")
    if resume_skipped > 0:
        print(f"[WARNING] Resumes skipped: {resume_skipped}")
    print()

    # Upload job descriptions
    print("UPLOADING JOB DESCRIPTIONS")
    print("-" * 30)
    job_uploaded, job_skipped = bulk_upload_job_descriptions()
    print(f"[SUCCESS] Job descriptions uploaded: {job_uploaded}")
    if job_skipped > 0:
        print(f"[WARNING] Job descriptions skipped: {job_skipped}")
    print()

    # Summary
    print("BULK UPLOAD COMPLETE")
    print("=" * 50)
    print(f"Total resumes processed: {resume_uploaded + resume_skipped}")
    print(f"Total job descriptions processed: {job_uploaded + job_skipped}")
    print(f"Total successful uploads: {resume_uploaded + job_uploaded}")
    if resume_skipped + job_skipped > 0:
        print(f"Total skipped: {resume_skipped + job_skipped}")

    if args.skip_ai:
        print("\n[INFO] Files uploaded without AI analysis.")
        print("       Run individual analysis later or re-run without --skip-ai")
    else:
        print("\n[INFO] Files are processed with AI analysis enabled.")
    print("       You can now access the web interface to view rankings and analytics.")

if __name__ == "__main__":
    main()
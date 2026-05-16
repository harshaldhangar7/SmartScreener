import random
import os
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Sample data for generating resumes - Indian names and context
FIRST_NAMES = [
    'Rahul', 'Priya', 'Amit', 'Sneha', 'Vikram', 'Anjali', 'Rajesh', 'Kavita',
    'Suresh', 'Meera', 'Arun', 'Pooja', 'Vivek', 'Kiran', 'Manoj', 'Sunita',
    'Ravi', 'Neha', 'Ajay', 'Divya', 'Sanjay', 'Rekha', 'Vinod', 'Anita',
    'Prakash', 'Sarika', 'Rakesh', 'Komal', 'Deepak', 'Preeti', 'Arjun', 'Kavya',
    'Rohan', 'Ananya', 'Karthik', 'Shreya', 'Aditya', 'Pallavi', 'Nikhil', 'Swati',
    'Siddharth', 'Aishwarya', 'Varun', 'Radhika', 'Abhishek', 'Nandini', 'Rohit', 'Sakshi',
    'Mohit', 'Tanvi', 'Gaurav', 'Shruti', 'Harsh', 'Anjali', 'Yash', 'Kritika',
    'Dhruv', 'Isha', 'Akash', 'Madhuri', 'Prateek', 'Simran', 'Utkarsh', 'Riya',
    'Aniket', 'Aaradhya', 'Dev', 'Trisha', 'Kartik', 'Bhavya', 'Saurabh', 'Nisha',
    'Aryan', 'Diya', 'Veer', 'Anvesha', 'Ishaan', 'Myra', 'Advait', 'Kiara'
]

LAST_NAMES = [
    'Sharma', 'Patel', 'Singh', 'Kumar', 'Gupta', 'Verma', 'Jain', 'Agarwal',
    'Yadav', 'Mishra', 'Chauhan', 'Pandey', 'Tiwari', 'Joshi', 'Nair', 'Iyer',
    'Menon', 'Pillai', 'Rao', 'Reddy', 'Naidu', 'Murthy', 'Sastri', 'Acharya',
    'Desai', 'Shah', 'Mehta', 'Trivedi', 'Kapoor', 'Malhotra', 'Khanna', 'Saxena',
    'Bhatia', 'Chopra', 'Kohli', 'Dhawan', 'Gill', 'Bansal', 'Arora', 'Sodhi',
    'Bakshi', 'Dutta', 'Ghosh', 'Banerjee', 'Chatterjee', 'Mukherjee', 'Sen', 'Das',
    'Roy', 'Chakraborty', 'Nair', 'Pillai', 'Nambiar', 'Kurian', 'Thomas', 'Philip',
    'Mathew', 'Jacob', 'Isaac', 'Samuel', 'Daniel', 'John', 'Peter', 'George',
    'Fernandes', 'D\'Souza', 'Rodrigues', 'Fernandes', 'Almeida', 'Pereira', 'Dias'
]

JOB_ROLES = [
    'Software Engineer', 'Data Scientist', 'Product Manager', 'DevOps Engineer',
    'Frontend Developer', 'Backend Developer', 'Full Stack Developer',
    'Machine Learning Engineer', 'AI Research Scientist', 'Cloud Architect',
    'Security Engineer', 'Database Administrator', 'System Administrator',
    'Mobile App Developer', 'QA Engineer', 'Technical Lead', 'Engineering Manager',
    'Solutions Architect', 'Data Engineer', 'Business Intelligence Analyst'
]

COMPANIES = [
    # Big MNCs in India
    'TCS', 'Infosys', 'Wipro', 'HCL Technologies', 'Tech Mahindra', 'Cognizant',
    'Accenture', 'IBM India', 'Microsoft India', 'Google India', 'Amazon India',
    'Oracle India', 'SAP India', 'Adobe India', 'Cisco India', 'Intel India',
    'Dell India', 'HP India', 'Lenovo India', 'Samsung India', 'Sony India',

    # Indian Unicorns & Major Startups
    'Flipkart', 'Paytm', 'Ola', 'Swiggy', 'Zomato', 'MakeMyTrip', 'BookMyShow',
    'Byju\'s', 'Unacademy', 'Cred', 'Razorpay', 'Freshworks', 'Zoho', 'BrowserStack',
    'Cure.fit', 'Nykaa', 'Meesho', 'Udaan', 'BigBasket', 'Grofers', 'Dunzo',
    'Rivigo', 'BlackBuck', 'Delhivery', 'Blue Dart', 'Shadowfax', 'Porter',

    # Indian Startups & Growing Companies
    'Myntra', 'Snapdeal', 'Jabong', 'ShopClues', 'Homeshop18', 'Naaptol',
    'Indiamart', 'Justdial', '99acres', 'MagicBricks', 'Housing.com', 'NoBroker',
    'Oyo', 'Treebo', 'FabHotels', 'Stayzilla', 'Goibibo', 'RedBus', 'AbhiBus',

    # Big Giants & Conglomerates
    'Reliance Industries', 'Tata Group', 'Birla Group', 'Adani Group', 'Mahindra Group',
    'Bajaj Group', 'ITC Limited', 'Hindustan Unilever', 'Nestle India', 'Coca-Cola India',
    'PepsiCo India', 'Procter & Gamble India', 'Johnson & Johnson India', 'Pfizer India'
]

TECHNOLOGIES = {
    'programming': ['Python', 'Java', 'JavaScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'TypeScript', 'Kotlin', 'Scala', 'Perl'],
    'web': ['React', 'Angular', 'Vue.js', 'Node.js', 'Express', 'Django', 'Flask', 'Spring Boot', 'Laravel', 'CodeIgniter', 'ASP.NET'],
    'data': ['SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Oracle', 'Redis', 'Elasticsearch', 'Kafka', 'Hadoop', 'Spark', 'Hive', 'Pig'],
    'cloud': ['AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI', 'Terraform', 'Ansible', 'OpenStack'],
    'ml': ['TensorFlow', 'PyTorch', 'Scikit-learn', 'Keras', 'Pandas', 'NumPy', 'NLTK', 'spaCy', 'OpenCV', 'Apache Mahout'],
    'mobile': ['React Native', 'Flutter', 'Android', 'iOS', 'Ionic', 'Cordova', 'Xamarin', 'PhoneGap']
}

UNIVERSITIES = [
    # IITs (Indian Institutes of Technology)
    'IIT Delhi', 'IIT Bombay', 'IIT Madras', 'IIT Kanpur', 'IIT Kharagpur',
    'IIT Roorkee', 'IIT Guwahati', 'IIT Hyderabad', 'IIT Indore', 'IIT Bhubaneswar',
    'IIT Patna', 'IIT Ropar', 'IIT Mandi', 'IIT Jodhpur', 'IIT Gandhinagar',
    'IIT Tirupati', 'IIT Palakkad', 'IIT Dharwad', 'IIT Jammu', 'IIT Bhilai',

    # NITs (National Institutes of Technology)
    'NIT Trichy', 'NIT Surathkal', 'NIT Warangal', 'NIT Calicut', 'NIT Rourkela',
    'NIT Jaipur', 'NIT Allahabad', 'NIT Durgapur', 'NIT Jamshedpur', 'NIT Nagpur',
    'NIT Kurukshetra', 'NIT Silchar', 'NIT Hamirpur', 'NIT Meghalaya', 'NIT Agartala',
    'NIT Goa', 'NIT Puducherry', 'NIT Manipur', 'NIT Srinagar', 'NIT Uttarakhand',

    # Other Premier Institutions
    'BITS Pilani', 'IIIT Hyderabad', 'IIIT Bangalore', 'IIIT Delhi', 'IIIT Allahabad',
    'Delhi University', 'JNU Delhi', 'Jamia Millia Islamia', 'Delhi Technological University',
    'Banaras Hindu University', 'Aligarh Muslim University', 'Jadavpur University',
    'Anna University', 'University of Madras', 'University of Mumbai', 'University of Pune',
    'Osmania University', 'Andhra University', 'Kakatiya University', 'Osmania University',
    'University of Hyderabad', 'Jawaharlal Nehru Technological University', 'SRM University',
    'VIT University', 'Amrita Vishwa Vidyapeetham', 'Manipal University', 'Christ University',
    'Mount Carmel College', 'St. Xavier\'s College', 'Loyola College', 'Presidency University'
]

DEGREES = [
    # Science Degrees
    'Bachelor of Science', 'Master of Science', 'Bachelor of Computer Science', 'Master of Computer Science',
    'Bachelor of Technology', 'Master of Technology', 'Bachelor of Engineering', 'Master of Engineering',
    'Bachelor of Science in Information Technology', 'Master of Science in Information Technology',

    # Arts Degrees
    'Bachelor of Arts', 'Master of Arts', 'Bachelor of Arts in English', 'Master of Arts in English',
    'Bachelor of Arts in Economics', 'Master of Arts in Economics', 'Bachelor of Arts in Psychology',
    'Master of Arts in Psychology', 'Bachelor of Arts in Sociology', 'Master of Arts in Sociology',
    'Bachelor of Arts in History', 'Master of Arts in History', 'Bachelor of Arts in Political Science',
    'Master of Arts in Political Science', 'Bachelor of Fine Arts', 'Master of Fine Arts',

    # Commerce Degrees
    'Bachelor of Commerce', 'Master of Commerce', 'Bachelor of Business Administration',
    'Master of Business Administration', 'Bachelor of Commerce in Accounting', 'Master of Commerce in Accounting',
    'Bachelor of Commerce in Finance', 'Master of Commerce in Finance', 'Chartered Accountancy',
    'Company Secretary', 'Cost and Management Accountancy',

    # Other Professional Degrees
    'Bachelor of Laws', 'Master of Laws', 'Bachelor of Education', 'Master of Education',
    'Bachelor of Pharmacy', 'Master of Pharmacy', 'Doctor of Philosophy', 'Doctor of Medicine',
    'Bachelor of Dental Surgery', 'Bachelor of Ayurvedic Medicine and Surgery'
]

def generate_person():
    """Generate a random person with basic info."""
    first_name = random.choice(FIRST_NAMES)
    last_name = random.choice(LAST_NAMES)
    name = f"{first_name} {last_name}"
    email = f"{first_name.lower()}.{last_name.lower()}@{random.choice(['gmail.com', 'yahoo.com', 'outlook.com', 'hotmail.com', 'rediffmail.com', 'live.com'])}"
    # Indian phone number format: +91 XXXXX XXXXX
    phone = f"+91 {random.randint(70000,99999)} {random.randint(10000,99999)}"

    # Generate experience years (0-15)
    experience_years = random.uniform(0, 15)

    return {
        'name': name,
        'email': email,
        'phone': phone,
        'experience_years': round(experience_years, 1)
    }

def generate_education():
    """Generate education history."""
    num_degrees = random.randint(1, 2)
    education = []

    for i in range(num_degrees):
        degree = random.choice(DEGREES)
        university = random.choice(UNIVERSITIES)
        field = random.choice(['Computer Science', 'Software Engineering', 'Data Science', 'Information Technology', 'Electrical Engineering'])
        year = datetime.now().year - random.randint(2, 8) - (i * 4)

        education.append({
            'degree': degree,
            'field': field,
            'university': university,
            'year': year
        })

    return education

def generate_experience(person, experience_years):
    """Generate work experience."""
    experience = []
    current_year = datetime.now().year
    remaining_years = experience_years

    while remaining_years > 0:
        job_role = random.choice(JOB_ROLES)
        company = random.choice(COMPANIES)

        # Duration between 1-4 years, but not exceeding remaining experience
        duration = min(random.uniform(1, 4), remaining_years)
        end_year = current_year
        start_year = int(end_year - duration)

        experience.append({
            'title': job_role,
            'company': company,
            'start_year': start_year,
            'end_year': end_year if duration >= 1 else None,
            'duration': round(duration, 1)
        })

        current_year = start_year
        remaining_years -= duration

    return experience

def generate_skills(job_role):
    """Generate relevant skills based on job role."""
    skills = []

    # Add role-specific skills
    if 'Data' in job_role or 'ML' in job_role or 'AI' in job_role:
        skills.extend(random.sample(TECHNOLOGIES['data'], random.randint(2, 4)))
        skills.extend(random.sample(TECHNOLOGIES['ml'], random.randint(2, 4)))
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(1, 3)))
    elif 'Frontend' in job_role or 'Web' in job_role:
        skills.extend(random.sample(TECHNOLOGIES['web'], random.randint(3, 5)))
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(1, 2)))
    elif 'Backend' in job_role:
        skills.extend(random.sample(TECHNOLOGIES['web'], random.randint(2, 4)))
        skills.extend(random.sample(TECHNOLOGIES['data'], random.randint(1, 3)))
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(2, 3)))
    elif 'DevOps' in job_role or 'Cloud' in job_role:
        skills.extend(random.sample(TECHNOLOGIES['cloud'], random.randint(3, 5)))
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(1, 2)))
    elif 'Mobile' in job_role:
        skills.extend(random.sample(TECHNOLOGIES['mobile'], random.randint(2, 4)))
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(1, 2)))
    else:
        # General software engineering
        skills.extend(random.sample(TECHNOLOGIES['programming'], random.randint(2, 4)))
        skills.extend(random.sample(TECHNOLOGIES['web'], random.randint(1, 3)))
        skills.extend(random.sample(TECHNOLOGIES['data'], random.randint(1, 2)))

    # Remove duplicates and limit to 8-12 skills
    skills = list(set(skills))[:random.randint(8, 12)]
    random.shuffle(skills)

    return skills

def generate_summary(person, skills, experience):
    """Generate a professional summary in Indian style."""
    exp_years = int(person['experience_years'])
    primary_skill = skills[0] if skills else 'software development'

    # Indian-style summaries with career objectives
    summaries = [
        f"A dynamic and result-oriented {primary_skill} professional with {exp_years}+ years of comprehensive experience in the IT industry. Seeking challenging opportunities to leverage technical expertise and contribute to organizational growth.",
        f"Highly motivated {primary_skill} engineer with {exp_years} years of hands-on experience in developing scalable solutions. Proficient in modern technologies with a proven track record of delivering high-quality software products.",
        f"Experienced {primary_skill} specialist with {exp_years} years of expertise in full-stack development. Passionate about creating innovative solutions and continuously learning emerging technologies to stay ahead in the competitive IT landscape.",
        f"Detail-oriented {primary_skill} professional with {exp_years} years of experience in enterprise-level application development. Strong analytical skills combined with technical proficiency to deliver robust and efficient software solutions.",
        f"Accomplished {primary_skill} engineer with {exp_years} years of progressive experience in software development lifecycle. Adept at working in agile environments and committed to maintaining high standards of code quality and best practices."
    ]

    return random.choice(summaries)

def generate_career_objective(job_role, skills):
    """Generate Indian-style career objective."""
    objectives = [
        f"To work in a challenging environment where I can utilize my {', '.join(skills[:3])} skills and {job_role.lower()} expertise to contribute effectively to organizational goals while continuously enhancing my professional knowledge.",
        f"Seeking a challenging position as {job_role} where I can leverage my technical skills in {', '.join(skills[:2])} and contribute to innovative projects while growing professionally in a dynamic organization.",
        f"To obtain a challenging position in {job_role} role where I can apply my knowledge of {', '.join(skills[:3])} to develop innovative solutions and work collaboratively with a talented team.",
        f"Aspiring to work as {job_role} in a reputed organization where I can utilize my {', '.join(skills[:2])} expertise and contribute to the company's success while enhancing my professional skills."
    ]
    return random.choice(objectives)

def generate_achievements():
    """Generate Indian-style achievements."""
    achievements = [
        "Successfully led a team of 5 developers in delivering a critical project ahead of schedule",
        "Reduced application response time by 40% through performance optimization",
        "Implemented CI/CD pipeline resulting in 60% faster deployment cycles",
        "Mentored 3 junior developers and improved team productivity by 25%",
        "Received 'Star Performer' award for Q4 2023",
        "Contributed to open source projects with 500+ GitHub stars",
        "Certified in AWS Solutions Architect and Azure DevOps",
        "Published 2 technical articles on Medium with 10K+ views",
        "Completed 50+ hours of professional development training",
        "Led digital transformation initiative saving ₹2L annually"
    ]
    return random.sample(achievements, random.randint(2, 4))

def generate_projects(job_role):
    """Generate project descriptions."""
    projects = [
        {
            "name": "E-Commerce Platform Redesign",
            "tech": ["React", "Node.js", "MongoDB", "AWS"],
            "description": "Led the complete redesign of a major e-commerce platform serving 100K+ users. Implemented microservices architecture and improved page load time by 50%.",
            "duration": "8 months"
        },
        {
            "name": "Financial Analytics Dashboard",
            "tech": ["Python", "Django", "PostgreSQL", "Chart.js"],
            "description": "Developed a comprehensive analytics dashboard for financial data visualization. Integrated with multiple APIs and implemented real-time data processing.",
            "duration": "6 months"
        },
        {
            "name": "Healthcare Management System",
            "tech": ["Java", "Spring Boot", "MySQL", "Angular"],
            "description": "Built a complete healthcare management system for hospitals with patient records, appointment scheduling, and billing modules.",
            "duration": "12 months"
        },
        {
            "name": "IoT Smart Home Solution",
            "tech": ["Python", "Raspberry Pi", "MQTT", "Firebase"],
            "description": "Designed and implemented an IoT-based smart home automation system with mobile app control and voice integration.",
            "duration": "4 months"
        }
    ]
    return random.sample(projects, random.randint(1, 3))

def create_pdf_resume_indian(person, education, experience, skills, summary, career_objective, achievements, projects, filename):
    """Create a PDF resume in Indian style."""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        alignment=1  # Center
    )

    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=6,
        textColor='darkblue'
    )

    content = []

    # Header
    content.append(Paragraph(person['name'], title_style))
    content.append(Paragraph(f"{person['email']} | {person['phone']}", styles['Normal']))
    content.append(Spacer(1, 12))

    # Career Objective
    content.append(Paragraph("Career Objective", section_style))
    content.append(Paragraph(career_objective, styles['Normal']))
    content.append(Spacer(1, 12))

    # Professional Summary
    content.append(Paragraph("Professional Summary", section_style))
    content.append(Paragraph(summary, styles['Normal']))
    content.append(Spacer(1, 12))

    # Skills
    content.append(Paragraph("Technical Skills", section_style))
    skills_text = ", ".join(skills)
    content.append(Paragraph(skills_text, styles['Normal']))
    content.append(Spacer(1, 12))

    # Experience
    content.append(Paragraph("Professional Experience", section_style))
    for exp in experience:
        duration = f"{exp['start_year']} - {'Present' if exp['end_year'] is None else exp['end_year']}"
        content.append(Paragraph(f"<b>{exp['title']}</b>", styles['Normal']))
        content.append(Paragraph(f"<i>{exp['company']}</i>", styles['Normal']))
        content.append(Paragraph(f"{duration} ({exp['duration']} years)", styles['Italic']))
        content.append(Spacer(1, 6))

    # Projects
    if projects:
        content.append(Paragraph("Key Projects", section_style))
        for project in projects:
            content.append(Paragraph(f"<b>{project['name']}</b>", styles['Normal']))
            content.append(Paragraph(f"<i>Technologies: {', '.join(project['tech'])}</i>", styles['Italic']))
            content.append(Paragraph(project['description'], styles['Normal']))
            content.append(Paragraph(f"<i>Duration: {project['duration']}</i>", styles['Italic']))
            content.append(Spacer(1, 6))

    # Achievements
    content.append(Paragraph("Achievements & Certifications", section_style))
    for achievement in achievements:
        content.append(Paragraph(f"• {achievement}", styles['Normal']))
    content.append(Spacer(1, 12))

    # Education
    content.append(Paragraph("Education", section_style))
    for edu in education:
        content.append(Paragraph(f"<b>{edu['degree']} in {edu['field']}</b>", styles['Normal']))
        content.append(Paragraph(f"{edu['university']}, {edu['year']}", styles['Normal']))
        content.append(Spacer(1, 6))

    doc.build(content)

def create_docx_resume_indian(person, education, experience, skills, summary, career_objective, achievements, projects, filename):
    """Create a DOCX resume in Indian style."""
    doc = Document()

    # Header
    title = doc.add_heading(person['name'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact = doc.add_paragraph()
    contact.add_run(f"{person['email']} | {person['phone']}")
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Career Objective
    doc.add_heading('Career Objective', level=1)
    doc.add_paragraph(career_objective)

    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(summary)

    # Technical Skills
    doc.add_heading('Technical Skills', level=1)
    skills_para = doc.add_paragraph()
    skills_para.add_run(", ".join(skills))

    # Professional Experience
    doc.add_heading('Professional Experience', level=1)
    for exp in experience:
        duration = f"{exp['start_year']} - {'Present' if exp['end_year'] is None else exp['end_year']}"
        p = doc.add_paragraph()
        p.add_run(f"{exp['title']}").bold = True
        p.add_run(f"\n{exp['company']}")
        p.add_run(f"\n{duration} ({exp['duration']} years)").italic = True
        doc.add_paragraph()  # Add spacing

    # Key Projects
    if projects:
        doc.add_heading('Key Projects', level=1)
        for project in projects:
            p = doc.add_paragraph()
            p.add_run(f"{project['name']}").bold = True
            p.add_run(f"\nTechnologies: {', '.join(project['tech'])}").italic = True
            doc.add_paragraph(project['description'])
            p_duration = doc.add_paragraph()
            p_duration.add_run(f"Duration: {project['duration']}").italic = True
            doc.add_paragraph()  # Add spacing

    # Achievements & Certifications
    doc.add_heading('Achievements & Certifications', level=1)
    for achievement in achievements:
        doc.add_paragraph(f"• {achievement}")

    # Education
    doc.add_heading('Education', level=1)
    for edu in education:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']} in {edu['field']}").bold = True
        p.add_run(f"\n{edu['university']}, {edu['year']}")

    doc.save(filename)

def generate_resume(resume_id):
    """Generate a single resume with Indian style."""
    # Generate person
    person = generate_person()

    # Generate education
    education = generate_education()

    # Generate experience based on years
    experience = generate_experience(person, person['experience_years'])

    # Determine primary job role from most recent experience
    primary_role = experience[0]['title'] if experience else random.choice(JOB_ROLES)

    # Generate skills based on role
    skills = generate_skills(primary_role)

    # Generate summary
    summary = generate_summary(person, skills, experience)

    # Generate career objective
    career_objective = generate_career_objective(primary_role, skills)

    # Generate achievements
    achievements = generate_achievements()

    # Generate projects
    projects = generate_projects(primary_role)

    # Create filename using candidate name (replace spaces with underscores)
    safe_name = person['name'].replace(' ', '_')

    # Determine file format (alternate between PDF and DOCX)
    if resume_id % 2 == 0:
        filename = f"resumes/{safe_name}.pdf"
        create_pdf_resume_indian(person, education, experience, skills, summary, career_objective, achievements, projects, filename)
    else:
        filename = f"resumes/{safe_name}.docx"
        create_docx_resume_indian(person, education, experience, skills, summary, career_objective, achievements, projects, filename)

    print(f"Generated {filename}")

def main():
    """Generate 100 resumes."""
    print("Starting resume generation...")

    # Create resumes directory if it doesn't exist
    os.makedirs('resumes', exist_ok=True)

    # Generate 100 resumes
    for i in range(1, 101):
        generate_resume(i)

    print("Resume generation completed!")
    print(f"Generated {len([f for f in os.listdir('resumes') if f.endswith(('.pdf', '.docx'))])} resume files")

if __name__ == "__main__":
    main()
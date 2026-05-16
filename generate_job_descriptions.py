import random
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Expanded job description templates with 10+ skills each
JOB_TEMPLATES = [
    {
        "title": "Senior Python Developer",
        "company": "TechCorp Solutions",
        "skills": ["Python", "Django", "PostgreSQL", "AWS", "Docker", "Git", "REST APIs", "JavaScript", "React", "Kubernetes", "Redis", "Celery", "FastAPI", "GraphQL", "Microservices"],
        "experience": "3-5 years",
        "education": "Bachelor's degree in Computer Science or related field",
    },
    {
        "title": "Data Scientist",
        "company": "AnalyticsPro",
        "skills": ["Python", "Machine Learning", "TensorFlow", "Pandas", "NumPy", "SQL", "Tableau", "Scikit-learn", "Jupyter", "PyTorch", "Spark", "Hadoop", "R", "Statistics", "Deep Learning"],
        "experience": "2-4 years",
        "education": "Master's degree in Data Science, Statistics, or related field",
    },
    {
        "title": "Full Stack Developer",
        "company": "InnovateTech",
        "skills": ["JavaScript", "React", "Node.js", "MongoDB", "Express", "HTML", "CSS", "TypeScript", "Git", "Vue.js", "Angular", "MySQL", "GraphQL", "Webpack", "Jest"],
        "experience": "3-6 years",
        "education": "Bachelor's degree in Computer Science or Information Technology",
    },
    {
        "title": "DevOps Engineer",
        "company": "CloudTech Systems",
        "skills": ["AWS", "Docker", "Kubernetes", "Terraform", "Jenkins", "Git", "Linux", "Python", "Bash", "Ansible", "Prometheus", "Grafana", "ELK Stack", "CI/CD", "Infrastructure as Code"],
        "experience": "4-7 years",
        "education": "Bachelor's degree in Computer Science or related field",
    },
    {
        "title": "Frontend Developer",
        "company": "UI/UX Solutions",
        "skills": ["JavaScript", "React", "TypeScript", "HTML", "CSS", "Sass", "Webpack", "Babel", "Jest", "Cypress", "Figma", "Adobe XD", "Responsive Design", "Accessibility", "Performance Optimization"],
        "experience": "2-5 years",
        "education": "Bachelor's degree in Computer Science or Design",
    },
    {
        "title": "Backend Developer",
        "company": "ServerSide Inc",
        "skills": ["Java", "Spring Boot", "Hibernate", "MySQL", "PostgreSQL", "REST APIs", "Microservices", "Docker", "Kubernetes", "Redis", "RabbitMQ", "JUnit", "Maven", "Git", "Agile"],
        "experience": "3-6 years",
        "education": "Bachelor's degree in Computer Science or Engineering",
    },
    {
        "title": "Machine Learning Engineer",
        "company": "AI Innovations",
        "skills": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Keras", "Pandas", "NumPy", "Jupyter", "MLflow", "Docker", "AWS", "Azure", "Computer Vision", "NLP", "Deep Learning"],
        "experience": "3-5 years",
        "education": "Master's degree in Computer Science, AI, or related field",
    },
    {
        "title": "Mobile App Developer",
        "company": "AppWorks Mobile",
        "skills": ["React Native", "Flutter", "Dart", "JavaScript", "TypeScript", "iOS", "Android", "Firebase", "SQLite", "REST APIs", "GraphQL", "Git", "Xcode", "Android Studio", "App Store"],
        "experience": "2-5 years",
        "education": "Bachelor's degree in Computer Science or related field",
    },
    {
        "title": "Cybersecurity Analyst",
        "company": "SecureNet Solutions",
        "skills": ["Network Security", "Firewalls", "SIEM", "Intrusion Detection", "Vulnerability Assessment", "Penetration Testing", "Cryptography", "Linux", "Python", "Bash", "Wireshark", "Metasploit", "NIST", "ISO 27001"],
        "experience": "3-6 years",
        "education": "Bachelor's degree in Cybersecurity or Computer Science",
    },
    {
        "title": "Cloud Architect",
        "company": "CloudMasters Inc",
        "skills": ["AWS", "Azure", "GCP", "Terraform", "CloudFormation", "Docker", "Kubernetes", "Microservices", "Serverless", "Lambda", "API Gateway", "Cloud Security", "Cost Optimization", "Monitoring", "Scalability"],
        "experience": "5-8 years",
        "education": "Bachelor's degree in Computer Science or related field",
    }
]

# Additional data for randomization
COMPANIES = [
    "TechCorp Solutions", "AnalyticsPro", "InnovateTech", "CloudTech Systems", "UI/UX Solutions",
    "ServerSide Inc", "AI Innovations", "AppWorks Mobile", "SecureNet Solutions", "CloudMasters Inc",
    "DataFlow Systems", "WebDev Pro", "CodeMasters", "TechSavvy Inc", "Digital Dynamics",
    "SmartSoft Solutions", "NextGen Tech", "FutureWorks", "TechHub Global", "InnovateLabs"
]

# More granular experience levels from 0.4 to 20+ years
EXPERIENCE_LEVELS = [
    "0.4-1 years", "1-2 years", "2-3 years", "3-4 years", "4-5 years", "5-6 years",
    "6-7 years", "7-8 years", "8-9 years", "9-10 years", "10-12 years", "12-15 years",
    "15-18 years", "18-20 years", "20+ years"
]

# Education degrees in Arts, Commerce, Science, Engineering and their branches
EDUCATION_REQUIREMENTS = [
    # Arts
    "Bachelor of Arts (BA) in English Literature",
    "Bachelor of Arts (BA) in History",
    "Bachelor of Arts (BA) in Psychology",
    "Bachelor of Arts (BA) in Sociology",
    "Bachelor of Arts (BA) in Political Science",
    "Master of Arts (MA) in English Literature",
    "Master of Arts (MA) in History",
    "Master of Arts (MA) in Psychology",

    # Commerce
    "Bachelor of Commerce (BCom) in Accounting",
    "Bachelor of Commerce (BCom) in Finance",
    "Bachelor of Commerce (BCom) in Business Administration",
    "Bachelor of Commerce (BCom) in Marketing",
    "Bachelor of Commerce (BCom) in Economics",
    "Master of Commerce (MCom) in Accounting",
    "Master of Commerce (MCom) in Finance",
    "Master of Commerce (MCom) in Business Administration",

    # Science
    "Bachelor of Science (BSc) in Mathematics",
    "Bachelor of Science (BSc) in Physics",
    "Bachelor of Science (BSc) in Chemistry",
    "Bachelor of Science (BSc) in Biology",
    "Bachelor of Science (BSc) in Computer Science",
    "Bachelor of Science (BSc) in Statistics",
    "Master of Science (MSc) in Mathematics",
    "Master of Science (MSc) in Physics",
    "Master of Science (MSc) in Chemistry",
    "Master of Science (MSc) in Computer Science",

    # Engineering
    "Bachelor of Engineering (BE) in Computer Engineering",
    "Bachelor of Engineering (BE) in Mechanical Engineering",
    "Bachelor of Engineering (BE) in Electrical Engineering",
    "Bachelor of Engineering (BE) in Civil Engineering",
    "Bachelor of Engineering (BE) in Chemical Engineering",
    "Bachelor of Engineering (BE) in Electronics Engineering",
    "Bachelor of Technology (BTech) in Computer Science",
    "Bachelor of Technology (BTech) in Information Technology",
    "Bachelor of Technology (BTech) in Mechanical Engineering",
    "Bachelor of Technology (BTech) in Electrical Engineering",
    "Master of Engineering (ME) in Computer Engineering",
    "Master of Engineering (ME) in Mechanical Engineering",
    "Master of Technology (MTech) in Computer Science",
    "Master of Technology (MTech) in Information Technology"
]

def create_pdf_job_description(job_data, filename):
    """Create a PDF job description."""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        alignment=1  # Center
    )

    company_style = ParagraphStyle(
        'Company',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=6,
        textColor='darkblue'
    )

    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=6,
        textColor='darkgreen'
    )

    content = []

    # Header
    content.append(Paragraph(job_data['title'], title_style))
    content.append(Paragraph(f"Company: {job_data['company']}", company_style))
    content.append(Spacer(1, 12))

    # Split description into paragraphs and add them
    paragraphs = [p.strip() for p in job_data['description'].split('\n') if p.strip()]

    for paragraph in paragraphs:
        if paragraph.startswith('Job Title:') or paragraph.startswith('Position:') or paragraph.startswith('Job Position:'):
            content.append(Paragraph(paragraph, styles['Heading2']))
        elif paragraph.startswith('Key Responsibilities:') or paragraph.startswith('Responsibilities:') or paragraph.startswith('Key Requirements:'):
            content.append(Paragraph(paragraph, section_style))
        elif paragraph.startswith('•'):
            content.append(Paragraph(paragraph, styles['Normal']))
        elif paragraph.startswith('Required Skills:') or paragraph.startswith('Technical Skills Required:'):
            content.append(Paragraph(paragraph, section_style))
        elif paragraph.startswith('Minimum Experience:') or paragraph.startswith('Experience Level:') or paragraph.startswith('Experience Required:'):
            content.append(Paragraph(paragraph, section_style))
        elif paragraph.startswith('Education:') or paragraph.startswith('Education Requirements:'):
            content.append(Paragraph(paragraph, section_style))
        else:
            content.append(Paragraph(paragraph, styles['Normal']))

        content.append(Spacer(1, 6))

    doc.build(content)

def create_docx_job_description(job_data, filename):
    """Create a DOCX job description."""
    doc = Document()

    # Header
    title = doc.add_heading(job_data['title'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    company_para = doc.add_paragraph()
    company_para.add_run(f"Company: {job_data['company']}").bold = True
    company_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Add spacing

    # Add description content
    paragraphs = [p.strip() for p in job_data['description'].split('\n') if p.strip()]

    for paragraph in paragraphs:
        if paragraph.startswith('Job Title:') or paragraph.startswith('Position:') or paragraph.startswith('Job Position:'):
            doc.add_heading(paragraph, level=1)
        elif paragraph.startswith('Key Responsibilities:') or paragraph.startswith('Responsibilities:') or paragraph.startswith('Key Requirements:'):
            doc.add_heading(paragraph, level=2)
        elif paragraph.startswith('•'):
            p = doc.add_paragraph(paragraph)
        elif paragraph.startswith('Required Skills:') or paragraph.startswith('Technical Skills Required:'):
            doc.add_heading(paragraph, level=2)
        elif paragraph.startswith('Minimum Experience:') or paragraph.startswith('Experience Level:') or paragraph.startswith('Experience Required:'):
            doc.add_heading(paragraph, level=2)
        elif paragraph.startswith('Education:') or paragraph.startswith('Education Requirements:'):
            doc.add_heading(paragraph, level=2)
        else:
            doc.add_paragraph(paragraph)

    doc.save(filename)

def generate_random_job_description(base_template, index):
    """Generate a randomized job description based on a template."""
    import random

    # Randomize company
    company = random.choice(COMPANIES)

    # Randomize experience level
    experience = random.choice(EXPERIENCE_LEVELS)

    # Randomize education - select 2-4 requirements to create duplicates/variety
    num_education_reqs = random.randint(2, 4)
    education_list = random.sample(EDUCATION_REQUIREMENTS, num_education_reqs)
    education = '; '.join(education_list)

    # Randomize skills - ensure at least 15 skills by adding more if needed
    base_skills = base_template['skills'].copy()
    additional_skills = [
        "Agile", "Scrum", "Kanban", "TDD", "BDD", "Unit Testing", "Integration Testing",
        "Code Review", "Documentation", "Technical Writing", "Mentoring", "Team Leadership",
        "Problem Solving", "Analytical Thinking", "Communication", "Collaboration",
        "Time Management", "Project Management", "Requirements Analysis", "System Design",
        "Leadership", "Strategic Planning", "Risk Management", "Quality Assurance", "DevOps",
        "Continuous Integration", "Continuous Deployment", "Version Control", "API Design",
        "Database Design", "System Architecture", "Performance Tuning", "Security Best Practices",
        "Cloud Computing", "Microservices", "Containerization", "Orchestration", "Monitoring",
        "Logging", "Troubleshooting", "Debugging", "Code Optimization", "Scalability", "Reliability"
    ]

    # Add random additional skills to reach at least 15 total
    while len(base_skills) < 15:
        skill = random.choice(additional_skills)
        if skill not in base_skills:
            base_skills.append(skill)

    # Shuffle skills to randomize order
    random.shuffle(base_skills)

    # Create job description text
    description = f"""
    Job Title: {base_template['title']}

    We are looking for an experienced {base_template['title']} to join our dynamic team at {company}. This role involves working with cutting-edge technologies and contributing to innovative projects.

    Key Responsibilities:
    • Design and develop high-quality software solutions
    • Collaborate with cross-functional teams
    • Implement best practices and coding standards
    • Participate in code reviews and technical discussions
    • Contribute to system architecture and design decisions
    • Ensure code quality and performance optimization
    • Stay updated with latest industry trends and technologies

    Required Skills:
    • {', '.join(base_skills)}

    Minimum Experience: {experience} of relevant experience

    Education: {education}

    Preferred Qualifications:
    • Experience with modern development methodologies
    • Strong problem-solving and analytical skills
    • Excellent communication and teamwork abilities
    • Passion for learning and professional growth
    • Experience with agile development practices
    """

    return {
        "title": base_template['title'],
        "company": company,
        "skills": base_skills,
        "experience": experience,
        "education": education,
        "description": description.strip()
    }

def generate_job_descriptions():
    """Generate 50+ sample job description files in PDF or DOC format (one format per job)."""
    print("Starting job description generation...")

    # Create job_descriptions directory if it doesn't exist
    import os
    os.makedirs('job_descriptions', exist_ok=True)

    generated_files = []
    target_count = 55  # Generate exactly 55 job descriptions (one format each)

    # Track used job titles to ensure uniqueness
    used_titles = set()

    job_index = 1
    while len(generated_files) < target_count:  # 1 file per job
        # Select random base template
        base_template = random.choice(JOB_TEMPLATES)

        # Generate randomized job description
        job_data = generate_random_job_description(base_template, job_index)

        # Ensure unique job title by adding variations
        original_title = job_data['title']
        title_variations = [
            f"{original_title}",
            f"Senior {original_title}",
            f"Lead {original_title}",
            f"Principal {original_title}",
            f"Junior {original_title}",
            f"{original_title} Specialist",
            f"{original_title} Expert",
            f"{original_title} Consultant"
        ]

        # Find an unused title variation
        selected_title = None
        for title_var in title_variations:
            if title_var not in used_titles:
                selected_title = title_var
                used_titles.add(title_var)
                break

        # If all variations used, add a number
        if not selected_title:
            selected_title = f"{original_title} {job_index}"
            used_titles.add(selected_title)

        # Update job data with unique title
        job_data['title'] = selected_title

        # Create filename with job title and experience
        safe_title = selected_title.replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '')
        safe_experience = job_data['experience'].replace(' ', '_').replace('-', '_').replace('+', 'plus')
        base_filename = f"{safe_title}_{safe_experience}_{job_index}"

        # Randomly choose format (PDF or DOC)
        file_format = random.choice(['pdf', 'doc'])

        if file_format == 'pdf':
            filename = f"job_descriptions/{base_filename}.pdf"
            create_pdf_job_description(job_data, filename)
        else:  # doc format
            filename = f"job_descriptions/{base_filename}.doc"
            create_docx_job_description(job_data, filename)

        generated_files.append(filename)
        print(f"Generated {filename}")

        job_index += 1

        # Stop if we've generated enough
        if job_index > target_count:
            break

    print("Job description generation completed!")
    print(f"Generated {len(generated_files)} job description files ({len(generated_files)} unique jobs):")
    print("Files generated:")

    for file in generated_files:
        print(f"  - {file}")

    print(f"\nTotal: {len(generated_files)} unique job descriptions")
    print(f"Total files: {len(generated_files)} (PDF or DOC - one format per job)")

    print("\nThese files can be uploaded through the job description upload feature.")
    print("Each job description includes:")
    print("  - 15+ technical and soft skills")
    print("  - Unique job titles (no duplicates)")
    print("  - Randomized company names")
    print("  - Experience from 0.4 to 20+ years")
    print("  - Education degrees in Arts, Commerce, Science, Engineering and branches")
    print("  - Filenames include job title and experience level")

if __name__ == "__main__":
    generate_job_descriptions()
import os
from dotenv import load_dotenv
import logging
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session
import uuid
from werkzeug.utils import secure_filename
from resume_parser import parse_resume
from baseline_ranker import rank_candidates_baseline
# Delay import of ranker to avoid dependency during migrations
try:
	from candidate_ranker import rank_candidates
	RANKING_AVAILABLE = True
except Exception as _e:
	RANKING_AVAILABLE = False
	rank_candidates = None

# Import AI analyzer
try:
	from ai_resume_analyzer import ai_analyzer
	AI_ANALYSIS_AVAILABLE = True
except Exception as _e:
	AI_ANALYSIS_AVAILABLE = False
	ai_analyzer = None
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.ext.declarative import declarative_base
from datetime import datetime
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s %(name)s %(message)s'
)
logger = logging.getLogger(__name__)

# Lazy import embedding model to allow migrations without extra deps
try:
    import time
    start_time = time.time()
    logger.info("Loading SentenceTransformer model")
    from sentence_transformers import SentenceTransformer
    embedding_model = SentenceTransformer(os.environ.get("EMBEDDING_MODEL", "all-MiniLM-L6-v2"))
    load_time = time.time() - start_time
    logger.info(f"SentenceTransformer loaded in {load_time:.2f} seconds")
except Exception as _e:
    SentenceTransformer = None
    embedding_model = None
    logger.warning("Failed to load SentenceTransformer")
import base64
import json
from werkzeug.security import generate_password_hash, check_password_hash
from email_validator import validate_email, EmailNotValidError
from functools import wraps
from sqlalchemy import text as sa_text

# Load environment variables from .env file
load_dotenv()


# Initialize Flask app
app = Flask(__name__)
# Ensure SESSION_SECRET is set
session_secret = os.environ.get("SESSION_SECRET")
if not session_secret:
    raise ValueError("SESSION_SECRET environment variable is required")
app.secret_key = session_secret

# Initialize Flask-Limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Configure database
# Ensure DATABASE_URL is set
database_url = os.environ.get("DATABASE_URL")
if not database_url:
    raise ValueError("DATABASE_URL environment variable is required")
app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# Initialize SQLAlchemy
db = SQLAlchemy(app)

# Initialize embedding model (small general-purpose; can be swapped)
# embedding_model = SentenceTransformer(os.environ.get("EMBEDDING_MODEL", "all-MiniLM-L6-v2"))

def encode_vector_to_text(vector):
    try:
        return base64.b64encode(json.dumps(vector).encode("utf-8")).decode("utf-8")
    except Exception:
        # Fallback simple json string if encoding fails
        return json.dumps(vector)

def decode_text_to_vector(text_value):
    if not text_value:
        return []
    try:
        return json.loads(base64.b64decode(text_value.encode("utf-8")).decode("utf-8"))
    except Exception:
        try:
            return json.loads(text_value)
        except Exception:
            return []

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Database models
class Candidate(db.Model):
    __tablename__ = 'candidates'

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100))
    phone = db.Column(db.String(20))
    gender = db.Column(db.String(20))  # For bias analysis
    ethnicity = db.Column(db.String(50))  # For bias analysis
    skills = db.Column(db.JSON)  # List of skills
    experience_years = db.Column(db.Float, default=0)
    education = db.Column(db.JSON)  # List of education objects
    experience_entries = db.Column(db.JSON)  # List of experience objects
    resume_filename = db.Column(db.String(255))
    embedding = db.Column(db.Text)
    parsed_date = db.Column(db.DateTime, default=datetime.now)
    
    # AI Analysis fields
    ai_analysis = db.Column(db.JSON)  # Comprehensive AI analysis results
    ai_summary = db.Column(db.Text)  # AI-generated professional summary
    ai_quality_score = db.Column(db.Float, default=0)  # Overall quality score from AI
    ai_analysis_date = db.Column(db.DateTime)  # When AI analysis was performed

    def __repr__(self):
        return f"<Candidate(id={self.id}, name='{self.name}')>"

    def to_dict(self):
        try:
            return {
                'id': self.id,
                'name': self.name,
                'email': self.email if self.email else '',
                'phone': self.phone if self.phone else '',
                'gender': self.gender if self.gender else '',
                'ethnicity': self.ethnicity if self.ethnicity else '',
                'skills': self.skills if self.skills else [],
                'education': self.education if self.education else [],
                'experience_entries': self.experience_entries if self.experience_entries else [],
                'experience_years': float(self.experience_years) if self.experience_years is not None else 0.0,
                'resume_filename': self.resume_filename,
                'embedding': self.embedding,
                'parsed_date': self.parsed_date.isoformat() if self.parsed_date else None,
                'ai_analysis': self.ai_analysis if self.ai_analysis else {},
                'ai_summary': self.ai_summary if self.ai_summary else '',
                'ai_quality_score': float(self.ai_quality_score) if self.ai_quality_score is not None else 0.0,
                'ai_analysis_date': self.ai_analysis_date.isoformat() if self.ai_analysis_date else None
            }
        except Exception as e:
            logger.error(f"Error converting candidate to dict: {str(e)}")
            return {
                'id': self.id,
                'name': self.name,
                'email': '',
                'phone': '',
                'gender': '',
                'ethnicity': '',
                'skills': [],
                'education': [],
                'experience_entries': [],
                'experience_years': 0.0,
                'resume_filename': self.resume_filename,
                'parsed_date': None,
                'ai_analysis': {},
                'ai_summary': '',
                'ai_quality_score': 0.0,
                'ai_analysis_date': None
            }

class JobDescription(db.Model):
    __tablename__ = 'job_descriptions'

    id = db.Column(db.Integer, primary_key=True)
    company_name = db.Column(db.String(100), nullable=True)
    title = db.Column(db.String(100), nullable=False)
    required_skills = db.Column(db.JSON)  # List of required skills
    min_experience = db.Column(db.Float, default=0)
    required_education = db.Column(db.JSON)  # List of required education levels
    created_date = db.Column(db.DateTime, default=datetime.now)
    fk_user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    embedding = db.Column(db.Text)

    def __repr__(self):
        return f"<JobDescription(id={self.id}, title='{self.title}')>"

    def to_dict(self):
        try:
            return {
                'id': self.id,
                'company_name': self.company_name if self.company_name else '',
                'title': self.title,
                'required_skills': self.required_skills if self.required_skills else [],
                'min_experience': float(self.min_experience) if self.min_experience is not None else 0.0,
                'required_education': self.required_education if self.required_education else [],
                'created_date': self.created_date.isoformat() if self.created_date else None,
                'embedding': self.embedding
            }
        except Exception as e:
            logger.error(f"Error converting job description to dict: {str(e)}")
            return {
                'id': self.id,
                'company_name': '',
                'title': self.title,
                'required_skills': [],
                'min_experience': 0.0,
                'required_education': [],
                'created_date': None
            }

# User model
class User(UserMixin, db.Model):
    __tablename__ = 'users'

    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(255))
    is_admin = db.Column(db.Boolean, default=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Create database tables
with app.app_context():
    import time
    db_start = time.time()
    logger.info("Creating database tables")
    db.create_all()
    db_time = time.time() - db_start
    logger.info(f"Database tables created in {db_time:.2f} seconds")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard with stats and workflow management."""
    import time
    start_time = time.time()
    logger.info("Loading dashboard")
    try:
        # Get statistics
        total_candidates = Candidate.query.count()
        total_jobs = JobDescription.query.filter_by(fk_user_id=current_user.id).count()
        ai_analyzed = Candidate.query.filter(Candidate.ai_analysis.isnot(None)).count()

        # Calculate average quality score
        candidates_with_scores = Candidate.query.filter(Candidate.ai_quality_score > 0).all()
        avg_quality_score = 0
        if candidates_with_scores:
            avg_quality_score = sum(c.ai_quality_score for c in candidates_with_scores) / len(candidates_with_scores)

        # Calculate average traditional quality score (heuristic based on skills and experience)
        avg_traditional_score = 0
        if total_candidates > 0:
            traditional_scores = []
            for candidate in Candidate.query.all():
                # Simple heuristic: skills count * 5 + experience * 3, capped at 100
                skills_score = len(candidate.skills or []) * 5
                exp_score = (candidate.experience_years or 0) * 3
                total_traditional = min(100, skills_score + exp_score)
                traditional_scores.append(total_traditional)
            if traditional_scores:
                avg_traditional_score = sum(traditional_scores) / len(traditional_scores)

        # Get recent activities (mock data for now)
        recent_activities = [
            {
                'type': 'upload',
                'icon': 'upload',
                'title': 'Resume Uploaded',
                'description': f'{total_candidates} candidates in database',
                'timestamp': 'Just now'
            },
            {
                'type': 'job',
                'icon': 'briefcase',
                'title': 'Job Created',
                'description': f'{total_jobs} active job descriptions',
                'timestamp': '2 hours ago'
            }
        ] if total_candidates > 0 or total_jobs > 0 else []

        # Get top candidates
        top_candidates = Candidate.query.order_by(Candidate.ai_quality_score.desc()).limit(3).all()

        # Quality distribution for chart
        quality_distribution = [0, 0, 0, 0]  # Excellent, Good, Average, Below Average
        for candidate in candidates_with_scores:
            score = candidate.ai_quality_score
            if score >= 80:
                quality_distribution[0] += 1
            elif score >= 60:
                quality_distribution[1] += 1
            elif score >= 40:
                quality_distribution[2] += 1
            else:
                quality_distribution[3] += 1

        stats = {
            'total_candidates': total_candidates,
            'total_jobs': total_jobs,
            'ai_analyzed': ai_analyzed,
            'avg_quality_score': round(avg_quality_score, 1),
            'avg_traditional_score': round(avg_traditional_score, 1),
            'new_candidates_today': 0,  # Would need to track this
            'jobs_with_candidates': total_jobs,  # Simplified
            'total_rankings': 0  # Would need to track this
        }

        end_time = time.time()
        logger.info(f"Dashboard loaded in {end_time - start_time:.2f} seconds")
        return render_template('dashboard.html',
                              stats=stats,
                              recent_activities=recent_activities,
                              top_candidates=[c.to_dict() for c in top_candidates],
                              quality_distribution=quality_distribution,
                              ai_available=AI_ANALYSIS_AVAILABLE and ai_analyzer and ai_analyzer.is_available())

    except Exception as e:
        logger.error(f"Dashboard error: {e}")
        flash('Error loading dashboard data', 'danger')
        return render_template('dashboard.html',
                              stats={'total_candidates': 0, 'total_jobs': 0, 'ai_analyzed': 0, 'avg_quality_score': 0},
                              recent_activities=[],
                              top_candidates=[],
                              quality_distribution=[0, 0, 0, 0],
                              ai_available=False)

@app.route('/upload')
@login_required
def upload_page():
    """Dedicated upload page with enhanced UI."""
    return render_template('index.html')

@app.route('/upload_resume', methods=['POST'])
@login_required
def upload_resume():
    """Handle resume upload, parsing and storage."""
    try:
        if 'resume' not in request.files:
            flash('No file uploaded', 'danger')
            return redirect(url_for('index'))

        consent = request.form.get('consent')
        if not consent:
            flash('You must consent to the privacy policy.', 'danger')
            return redirect(url_for('index'))

        file = request.files['resume']

        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(url_for('index'))

        if not allowed_file(file.filename):
            flash('Invalid file type. Please upload a PDF or DOCX file.', 'danger')
            return redirect(url_for('index'))

        # Generate a unique filename to avoid collisions
        original_filename = secure_filename(file.filename)
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)

        try:
            file.save(filepath)
            
            # Extract raw text for AI analysis
            raw_text = ""
            try:
                if file_extension == '.pdf':
                    import fitz
                    doc = fitz.open(filepath)
                    for page in doc:
                        raw_text += page.get_text()
                elif file_extension in ['.docx', '.doc']:
                    import docx
                    doc = docx.Document(filepath)
                    for paragraph in doc.paragraphs:
                        raw_text += paragraph.text + '\n'
            except Exception as e:
                logger.warning(f"Failed to extract raw text for AI analysis: {e}")
            
            # Parse the resume
            candidate_data = parse_resume(filepath, original_filename)

            # Validate parsed data
            if not candidate_data.get('name'):
                raise ValueError("Could not extract candidate name from resume")

            # Create new candidate from parsed data
            new_candidate = Candidate(
                name=candidate_data.get('name', ''),
                email=candidate_data.get('email', ''),
                phone=candidate_data.get('phone', ''),
                skills=candidate_data.get('skills', []),
                experience_years=float(candidate_data.get('experience_years', 0)),
                education=candidate_data.get('education', []),
                experience_entries=candidate_data.get('experience_entries', []),
                resume_filename=original_filename,
                parsed_date=datetime.now()
            )

            # Save candidate data to database
            db.session.add(new_candidate)
            db.session.commit()

            # Check analysis type
            analysis_type = request.form.get('analysis_type', 'ai')

            # Perform AI analysis if requested and available
            if analysis_type == 'ai' and AI_ANALYSIS_AVAILABLE and ai_analyzer and ai_analyzer.is_available():
                try:
                    logger.info(f"Performing AI analysis for candidate: {new_candidate.name}")
                    ai_analysis = ai_analyzer.analyze_resume_comprehensive(raw_text, candidate_data)
                    ai_summary = ai_analyzer.generate_candidate_summary(candidate_data, ai_analysis)

                    # Update candidate with AI analysis
                    new_candidate.ai_analysis = ai_analysis
                    new_candidate.ai_summary = ai_summary
                    new_candidate.ai_quality_score = ai_analysis.get('overall_quality_score', 0)
                    new_candidate.ai_analysis_date = datetime.now()

                    logger.info(f"AI analysis completed for {new_candidate.name} with quality score: {new_candidate.ai_quality_score}")
                except Exception as e:
                    logger.error(f"AI analysis failed for {new_candidate.name}: {e}")
            elif analysis_type == 'traditional':
                logger.info(f"Skipping AI analysis for {new_candidate.name} - traditional analysis selected")
                # For traditional analysis, we could compute a basic quality score
                # Using the same heuristic as the dashboard
                skills_score = len(candidate_data.get('skills', [])) * 5
                exp_score = float(candidate_data.get('experience_years', 0)) * 3
                traditional_quality = min(100, skills_score + exp_score)
                new_candidate.ai_quality_score = traditional_quality  # Store in same field for consistency
                logger.info(f"Traditional quality score computed for {new_candidate.name}: {traditional_quality}")

            # Compute and store candidate embedding (post-commit to get ID if needed)
            try:
                if embedding_model is not None:
                    text_parts = [
                        new_candidate.name or '',
                        ', '.join(new_candidate.skills or []),
                        ', '.join([exp.get('title','') for exp in (new_candidate.experience_entries or []) if isinstance(exp, dict)]),
                        ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (new_candidate.education or [])])
                    ]
                    text_value = ' \n '.join([p for p in text_parts if p])
                    vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                    new_candidate.embedding = encode_vector_to_text(vector)
                    db.session.commit()
            except Exception as e:
                logger.error(f"Failed to compute candidate embedding: {str(e)}")

            success_message = 'Resume uploaded and parsed successfully!'
            if analysis_type == 'ai':
                if AI_ANALYSIS_AVAILABLE and ai_analyzer and ai_analyzer.is_available():
                    success_message += ' AI analysis completed.'
                else:
                    success_message += ' (AI analysis not available - configure OpenAI API key for enhanced features)'
            else:
                success_message += ' Traditional analysis completed.'
            
            flash(success_message, 'success')
            return redirect(url_for('candidates'))

        except ValueError as ve:
            flash(f'Error processing resume: {str(ve)}', 'danger')
            return redirect(url_for('index'))
        except Exception as e:
            logger.error(f"Error processing resume: {str(e)}")
            flash('An unexpected error occurred while processing the resume. Please try again.', 'danger')
            return redirect(url_for('index'))
        finally:
            # Clean up the file
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.error(f"Error removing temporary file {filepath}: {str(e)}")

    except Exception as e:
        logger.error(f"Error in upload_resume: {str(e)}")
        flash('An unexpected error occurred. Please try again.', 'danger')
        return redirect(url_for('index'))

@app.route('/candidates')
@login_required
def candidates():
    import time
    start_time = time.time()
    logger.info("Loading candidates page")
    candidates = Candidate.query.all()
    candidate_dicts = [c.to_dict() for c in candidates]
    end_time = time.time()
    logger.info(f"Candidates page loaded in {end_time - start_time:.2f} seconds for {len(candidate_dicts)} candidates")
    return render_template('candidates.html', candidates=candidate_dicts)

@app.route('/candidate/<int:candidate_id>')
def candidate_details(candidate_id):
    """Display details for a specific candidate."""
    candidate = Candidate.query.get(candidate_id)
    if candidate:
        return jsonify(candidate.to_dict())
    return jsonify({"error": "Candidate not found"}), 404

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/candidate/<int:candidate_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_candidate(candidate_id):
    """Delete a candidate. Only admin can delete."""
    candidate = Candidate.query.get(candidate_id)
    if not candidate:
        flash('Candidate not found', 'danger')
    else:
        db.session.delete(candidate)
        db.session.commit()
        flash('Candidate deleted successfully!', 'success')
    return redirect(url_for('candidates'))

@app.route('/candidate/<int:candidate_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_candidate(candidate_id):
    """Edit a candidate's information. Both admin and users can edit."""
    candidate = Candidate.query.get(candidate_id)
    if not candidate:
        flash('Candidate not found', 'danger')
        return redirect(url_for('candidates'))

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        # Validate name
        if not name:
            flash('Candidate name is required.', 'danger')
            return redirect(url_for('edit_candidate', candidate_id=candidate_id))
        
        candidate.name = name
        candidate.email = request.form.get('email', '').strip()
        candidate.phone = request.form.get('phone', '').strip()
        skills_text = request.form.get('skills', '')
        candidate.skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
        try:
            candidate.experience_years = float(request.form.get('experience_years', 0))
        except ValueError:
            candidate.experience_years = 0
        # Update education entries
        degrees = request.form.getlist('education_degrees[]')
        fields = request.form.getlist('education_fields[]')
        universities = request.form.getlist('education_universities[]')
        years = request.form.getlist('education_years[]')
        updated_education = []
        for idx in range(max(len(degrees), len(fields), len(universities), len(years))):
            deg = (degrees[idx] if idx < len(degrees) else '').strip()
            fld = (fields[idx] if idx < len(fields) else '').strip()
            uni = (universities[idx] if idx < len(universities) else '').strip()
            yr = (years[idx] if idx < len(years) else '').strip()
            if deg or fld or uni or yr:
                updated_education.append({
                    'degree': deg,
                    'field': fld,
                    'university': uni,
                    'year': yr
                })
        if updated_education:
            candidate.education = updated_education
        db.session.commit()
        
        # Recompute candidate embedding
        try:
            if embedding_model is not None:
                text_parts = [
                    candidate.name or '',
                    ', '.join(candidate.skills or []),
                    ', '.join([exp.get('title','') for exp in (candidate.experience_entries or []) if isinstance(exp, dict)]),
                    ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (candidate.education or [])])
                ]
                text_value = ' \n '.join([p for p in text_parts if p])
                vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                candidate.embedding = encode_vector_to_text(vector)
                db.session.commit()
        except Exception as e:
            logger.error(f"Failed to recompute candidate embedding: {str(e)}")
        flash('Candidate updated successfully!', 'success')
        return redirect(url_for('candidates'))
    return render_template('edit_candidate.html', candidate=candidate.to_dict())

@app.route('/job_descriptions')
@login_required
def job_descriptions():
    """List all job descriptions for the current user."""
    job_descriptions = JobDescription.query.filter_by(fk_user_id=current_user.id).all()
    return render_template('job_description.html', job_descriptions=[job.to_dict() for job in job_descriptions])

@app.route('/job_description/add', methods=['GET', 'POST'])
@login_required
def add_job_description():
    """Add a new job description."""
    if request.method == 'POST':
        company_name = request.form.get('company_name', '').strip()
        title = request.form.get('title', '').strip()
        # Validate title
        if not title:
            flash('Job title is required.', 'danger')
            return redirect(url_for('add_job_description'))

        skills_text = request.form.get('required_skills', '')
        required_skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]

        # Validate skills
        if not required_skills:
            flash('At least one required skill is needed.', 'danger')
            return redirect(url_for('add_job_description'))

        try:
            min_experience = float(request.form.get('min_experience', 0))
        except ValueError:
            min_experience = 0

        # Handle education requirements
        education_types = request.form.getlist('education_types[]')
        education_levels = request.form.getlist('education_levels[]')
        education_fields = request.form.getlist('education_fields[]')
        required_education = []

        for edu_type, level, field in zip(education_types, education_levels, education_fields):
            if level:  # Only add if level is selected
                edu_requirement = {
                    'type': edu_type.strip() if edu_type else '',
                    'degree': level,  # Keep 'degree' for backward compatibility, but it now represents level
                    'field': field.strip() if field else ''
                }
                required_education.append(edu_requirement)

        new_job = JobDescription(
            company_name=company_name,
            title=title,
            required_skills=required_skills,
            min_experience=min_experience,
            required_education=required_education,
            fk_user_id=current_user.id
        )
        db.session.add(new_job)
        db.session.commit()
        # Compute and store job embedding
        try:
            if embedding_model is not None:
                text_parts = [
                    company_name or '',
                    title or '',
                    ', '.join(required_skills or []),
                    ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (required_education or [])])
                ]
                text_value = ' \n '.join([p for p in text_parts if p])
                vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                new_job.embedding = encode_vector_to_text(vector)
                db.session.commit()
        except Exception as e:
            logger.error(f"Failed to compute job embedding: {str(e)}")
        flash('Job description added successfully!', 'success')
        return redirect(url_for('job_descriptions'))
    return render_template('add_job_description.html')

@app.route('/job_description/<int:job_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_job_description(job_id):
    """Edit a job description. Users can only edit their own job descriptions."""
    job = JobDescription.query.get_or_404(job_id)
    
    # Check if the user owns this job description
    if job.fk_user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to edit this job description.', 'danger')
        return redirect(url_for('job_descriptions'))

    if request.method == 'POST':
        company_name = request.form.get('company_name', '').strip()
        title = request.form.get('title', '').strip()
        # Validate title
        if not title:
            flash('Job title is required.', 'danger')
            return redirect(url_for('edit_job_description', job_id=job_id))

        job.company_name = company_name
        job.title = title
        skills_text = request.form.get('required_skills', '')
        job.required_skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]

        # Validate skills
        if not job.required_skills:
            flash('At least one required skill is needed.', 'danger')
            return redirect(url_for('edit_job_description', job_id=job_id))

        try:
            job.min_experience = float(request.form.get('min_experience', 0))
        except ValueError:
            job.min_experience = 0

        education_types = request.form.getlist('education_types[]')
        education_levels = request.form.getlist('education_levels[]')
        education_fields = request.form.getlist('education_fields[]')
        required_education = []
        for edu_type, level, field in zip(education_types, education_levels, education_fields):
            if level:
                edu_requirement = {
                    'type': edu_type.strip() if edu_type else '',
                    'degree': level,  # Keep 'degree' for backward compatibility
                    'field': field.strip() if field else ''
                }
                required_education.append(edu_requirement)
        job.required_education = required_education
        db.session.commit()
        # Recompute embedding after edits
        try:
            if embedding_model is not None:
                text_parts = [
                    job.company_name or '',
                    job.title or '',
                    ', '.join(job.required_skills or []),
                    ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (job.required_education or [])])
                ]
                text_value = ' \n '.join([p for p in text_parts if p])
                vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                job.embedding = encode_vector_to_text(vector)
                db.session.commit()
        except Exception as e:
            logger.error(f"Failed to recompute job embedding: {str(e)}")
        flash('Job description updated successfully!', 'success')
        return redirect(url_for('job_descriptions'))
    return render_template('edit_job_description.html', job=job.to_dict())

@app.route('/job_description/<int:job_id>/delete', methods=['POST'])
@login_required
def delete_job_description(job_id):
    """Delete a job description. Users can only delete their own job descriptions."""
    job = JobDescription.query.get_or_404(job_id)

    # Check if the user owns this job description
    if job.fk_user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to delete this job description.', 'danger')
        return redirect(url_for('job_descriptions'))

    db.session.delete(job)
    db.session.commit()
    flash('Job description deleted successfully!', 'success')
    return redirect(url_for('job_descriptions'))

@app.route('/upload_job_description', methods=['POST'])
@login_required
def upload_job_description():
    """Handle job description file upload and processing."""
    try:
        if 'job_file' not in request.files:
            flash('No file uploaded', 'danger')
            return redirect(url_for('add_job_description'))

        file = request.files['job_file']

        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(url_for('add_job_description'))

        if not allowed_file(file.filename):
            flash('Invalid file type. Please upload a PDF, DOC, or DOCX file.', 'danger')
            return redirect(url_for('add_job_description'))

        # Generate a unique filename to avoid collisions
        original_filename = secure_filename(file.filename)
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)

        try:
            file.save(filepath)

            # Extract text from the job description file
            job_text = ""
            try:
                if file_extension == 'pdf':
                    import fitz
                    doc = fitz.open(filepath)
                    for page in doc:
                        job_text += page.get_text()
                elif file_extension in ['docx', 'doc']:
                    import docx
                    doc = docx.Document(filepath)
                    for paragraph in doc.paragraphs:
                        job_text += paragraph.text + '\n'
            except Exception as e:
                logger.warning(f"Failed to extract text from job description file: {e}")

            # Parse the job description
            job_data = parse_job_description(job_text, original_filename)

            # If auto-extract is enabled, save directly
            if request.form.get('auto_extract'):
                new_job = JobDescription(
                    company_name=job_data.get('company_name', ''),
                    title=job_data.get('title', f'Job Description from {original_filename}'),
                    required_skills=job_data.get('required_skills', []),
                    min_experience=float(job_data.get('min_experience', 0)),
                    required_education=job_data.get('required_education', []),
                    fk_user_id=current_user.id
                )
                db.session.add(new_job)
                db.session.commit()

                # Compute and store job embedding
                try:
                    if embedding_model is not None:
                        text_parts = [
                            new_job.company_name or '',
                            new_job.title or '',
                            ', '.join(new_job.required_skills or []),
                            ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (new_job.required_education or [])])
                        ]
                        text_value = ' \n '.join([p for p in text_parts if p])
                        vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                        new_job.embedding = encode_vector_to_text(vector)
                        db.session.commit()
                except Exception as e:
                    logger.error(f"Failed to compute job embedding: {str(e)}")

                flash('Job description uploaded and processed successfully!', 'success')
                return redirect(url_for('job_descriptions'))
            else:
                # Store extracted data in session for manual review
                session['extracted_job_data'] = job_data
                session['job_file_path'] = filepath
                return redirect(url_for('review_job_description'))

        except Exception as e:
            logger.error(f"Error processing job description file: {str(e)}")
            flash('An unexpected error occurred while processing the job description. Please try again.', 'danger')
            return redirect(url_for('add_job_description'))
        finally:
            # Clean up the file if it exists and we're not storing it for review
            if 'job_file_path' not in session and os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.error(f"Error removing temporary file {filepath}: {str(e)}")

    except Exception as e:
        logger.error(f"Error in upload_job_description: {str(e)}")
        flash('An unexpected error occurred. Please try again.', 'danger')
        return redirect(url_for('add_job_description'))

def parse_job_description(text, filename):
    """Parse job description text to extract structured data."""
    job_data = {
        'company_name': '',
        'title': 'Unknown Position',
        'required_skills': [],
        'min_experience': 0,
        'required_education': []
    }

    try:
        # Simple text processing to extract information
        text_lower = text.lower()

        # Extract company name (look for common patterns)
        company_patterns = [
            r'(?:company|organization|employer)[\s:]+([^\n\r]+)',
            r'(?:at|@)\s+([A-Z][A-Za-z\s]+(?:Inc|LLC|Corp|Corporation|Solutions|Ltd|Limited|Technologies|Systems|Labs)?)',
            r'([A-Z][A-Za-z\s]+(?:Inc|LLC|Corp|Corporation|Solutions|Ltd|Limited|Technologies|Systems|Labs))',
            r'(?:we are|join)\s+([A-Z][A-Za-z\s]+(?:Inc|LLC|Corp|Corporation|Solutions|Ltd|Limited|Technologies|Systems|Labs)?)'
        ]

        import re

        for i, pattern in enumerate(company_patterns):
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company_name = match.group(1).strip()
                # Clean up common suffixes and extra spaces
                company_name = re.sub(r'\s+', ' ', company_name)
                job_data['company_name'] = company_name
                logger.info(f"Found company name using pattern {i}: '{company_name}'")
                break
            else:
                pass

        if not job_data['company_name']:
            logger.warning("No company name found in job description")

        # Extract job title (look for common patterns)
        title_patterns = [
            r'(?:job title|position|role)[\s:]+([^\n\r]+)',
            r'^([^\n\r]+)(?:\n|\r|$)',
            r'(?:we are hiring|opening for)[\s:]+([^\n\r]+)'
        ]

        for pattern in title_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                job_data['title'] = match.group(1).strip()
                break

        # Extract skills (look for common skill keywords)
        skill_keywords = [
            'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 'typescript',
            'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring', 'sql', 'mysql',
            'postgresql', 'mongodb', 'aws', 'azure', 'docker', 'kubernetes', 'git'
        ]

        found_skills = []
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill.title())

        job_data['required_skills'] = found_skills[:10]  # Limit to 10 skills

        # Extract experience requirements
        exp_patterns = [
            r'(\d+)[\s\-]*(\d+)?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(?:minimum|min)\s*(\d+)[\s\-]*(\d+)?\s*(?:years?|yrs?)',
            r'experience[\s:]+(\d+)[\s\-]*(\d+)?\s*(?:years?|yrs?)',
            r'(?:experience\s*(?:level|required)[\s:]+)(\d+)[\s\-]*(\d+)?\s*(?:years?|yrs?)'
        ]

        for pattern in exp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                if match.group(2):
                    # Range like "2-5 years"
                    job_data['min_experience'] = float(match.group(1))
                else:
                    job_data['min_experience'] = float(match.group(1))
                break

        # Extract education requirements with smarter parsing
        # First, try to find complete education patterns
        education_patterns = [
            r'(bachelor|master|phd|doctorate)(?:\s+of)?\s+(science|engineering|arts|commerce|business|technology)(?:\s*\([^)]*\))?\s+in\s+([^\n,.]+)',
            r'(bsc|msc|phd|mba|btech|be|mtech|bca|bcom|ba)(?:\s*\([^)]*\))?\s+in\s+([^\n,.]+)',
            r'(bachelor|master|phd)(?:\s+degree)?(?:\s*\([^)]*\))?\s+in\s+([^\n,.]+)',
            r'(bachelor|master|phd|doctorate)(?:\s+of)?\s+(science|engineering|arts|commerce|business|technology)(?:\s*\([^)]*\))?\s*[:-]?\s*([^\n,.]+)',
        ]

        education_requirements = []

        for pattern in education_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if len(match) == 3:  # Pattern with level, type, and field
                    level, degree_type, field = match
                    # Map degree level
                    level_mapping = {
                        'bachelor': 'bachelor',
                        'bsc': 'bachelor',
                        'bca': 'bachelor',
                        'btech': 'bachelor',
                        'be': 'bachelor',
                        'ba': 'bachelor',
                        'bcom': 'bachelor',
                        'master': 'master',
                        'msc': 'master',
                        'mtech': 'master',
                        'mba': 'master',
                        'phd': 'phd',
                        'doctorate': 'phd'
                    }
                    mapped_level = level_mapping.get(level.lower(), 'bachelor')

                    # Map degree type
                    type_mapping = {
                        'science': 'science',
                        'engineering': 'engineering',
                        'arts': 'arts',
                        'commerce': 'commerce',
                        'business': 'business',
                        'technology': 'technology'
                    }
                    mapped_type = type_mapping.get(degree_type.lower(), 'science')

                    # Clean up field - don't override with generic mappings
                    field = field.strip().title()

                    education_requirements.append({
                        'type': mapped_type,
                        'degree': mapped_level,
                        'field': field
                    })
                    break  # Only take the first match from this pattern
            if education_requirements:
                break  # If we found education from patterns, stop looking

        # Fallback to keyword-based extraction if no patterns matched
        if not education_requirements:
            # Extract education requirements with type, level, and field
            education_keywords = {
                'phd': {'level': 'phd', 'type': 'science'},
                'doctorate': {'level': 'phd', 'type': 'science'},
                'master': {'level': 'master', 'type': 'science'},
                'msc': {'level': 'master', 'type': 'science'},
                'ms': {'level': 'master', 'type': 'science'},
                'mtech': {'level': 'master', 'type': 'engineering'},
                'mba': {'level': 'master', 'type': 'business'},
                'bachelor': {'level': 'bachelor', 'type': 'science'},
                'bsc': {'level': 'bachelor', 'type': 'science'},
                'bs': {'level': 'bachelor', 'type': 'science'},
                'btech': {'level': 'bachelor', 'type': 'engineering'},
                'be': {'level': 'bachelor', 'type': 'engineering'},
                'bca': {'level': 'bachelor', 'type': 'technology'},
                'bcom': {'level': 'bachelor', 'type': 'commerce'},
                'ba': {'level': 'bachelor', 'type': 'arts'},
                'engineering': {'level': 'bachelor', 'type': 'engineering'},
                'diploma': {'level': 'diploma', 'type': 'technology'},
                'certificate': {'level': 'certificate', 'type': 'technology'}
            }

            # Extract degree type keywords
            degree_types = {
                'science': ['computer science', 'physics', 'chemistry', 'biology', 'mathematics', 'statistics'],
                'engineering': ['engineering', 'mechanical', 'electrical', 'civil', 'chemical', 'software'],
                'arts': ['arts', 'humanities', 'literature', 'history', 'philosophy'],
                'commerce': ['commerce', 'business', 'finance', 'accounting', 'economics'],
                'business': ['business', 'management', 'marketing', 'entrepreneurship'],
                'technology': ['technology', 'information technology', 'computer', 'programming']
            }

            seen_combinations = set()  # Track unique combinations

            for keyword, info in education_keywords.items():
                if keyword in text_lower:
                    # Try to determine field from context - look for specific field names first
                    field = 'Computer Science'  # Default

                    # Look for specific field names in the text
                    specific_fields = ['computer science', 'information technology', 'software engineering',
                                     'mechanical engineering', 'electrical engineering', 'civil engineering',
                                     'business administration', 'finance', 'accounting', 'marketing']

                    for specific_field in specific_fields:
                        if specific_field in text_lower:
                            field = specific_field.title()
                            break

                    # If no specific field found, use generic mapping based on degree type
                    if field == 'Computer Science':
                        if info['type'] == 'engineering':
                            field = 'Engineering'
                        elif info['type'] == 'business':
                            field = 'Business Administration'
                        elif info['type'] == 'commerce':
                            field = 'Commerce'
                        elif info['type'] == 'arts':
                            field = 'Arts'
                        elif info['type'] == 'technology':
                            field = 'Information Technology'
                        # Keep 'Computer Science' as default for science type

                    # Create a unique key for this combination
                    combo_key = (info['type'], info['level'], field)

                    # Only add if we haven't seen this combination before
                    if combo_key not in seen_combinations:
                        seen_combinations.add(combo_key)
                        education_requirements.append({
                            'type': info['type'],
                            'degree': info['level'],
                            'field': field
                        })

        # If no specific education found, add a default one
        if not education_requirements:
            education_requirements.append({
                'type': 'science',
                'degree': 'master',
                'field': 'Computer Science'
            })

        job_data['required_education'] = education_requirements[:1]  # Limit to 1 for precision

    except Exception as e:
        logger.error(f"Error parsing job description: {e}")

    return job_data

@app.route('/review_job_description')
@login_required
def review_job_description():
    """Review and edit extracted job description data."""
    if 'extracted_job_data' not in session:
        flash('No job description data to review', 'warning')
        return redirect(url_for('add_job_description'))

    job_data = session['extracted_job_data']
    return render_template('review_job_description.html', job_data=job_data)

@app.route('/save_reviewed_job_description', methods=['POST'])
@login_required
def save_reviewed_job_description():
    """Save the reviewed and edited job description."""
    try:
        company_name = request.form.get('company_name', '').strip()
        title = request.form.get('title', '').strip()
        if not title:
            flash('Job title is required.', 'danger')
            return redirect(url_for('review_job_description'))

        skills_text = request.form.get('required_skills', '')
        required_skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]

        if not required_skills:
            flash('At least one required skill is needed.', 'danger')
            return redirect(url_for('review_job_description'))

        try:
            min_experience = float(request.form.get('min_experience', 0))
        except ValueError:
            min_experience = 0

        # Handle education requirements
        education_types = request.form.getlist('education_types[]')
        education_levels = request.form.getlist('education_levels[]')
        education_fields = request.form.getlist('education_fields[]')
        required_education = []

        for edu_type, level, field in zip(education_types, education_levels, education_fields):
            if level:  # Only add if level is selected
                edu_requirement = {
                    'type': edu_type.strip() if edu_type else '',
                    'degree': level,  # Keep 'degree' for backward compatibility
                    'field': field.strip() if field else ''
                }
                required_education.append(edu_requirement)

        new_job = JobDescription(
            company_name=company_name,
            title=title,
            required_skills=required_skills,
            min_experience=min_experience,
            required_education=required_education,
            fk_user_id=current_user.id
        )
        db.session.add(new_job)
        db.session.commit()

        # Compute and store job embedding
        try:
            if embedding_model is not None:
                text_parts = [
                    new_job.title or '',
                    ', '.join(new_job.required_skills or []),
                    ', '.join([f"{e.get('degree','')} {e.get('field','')}" if isinstance(e, dict) else str(e) for e in (new_job.required_education or [])])
                ]
                text_value = ' \n '.join([p for p in text_parts if p])
                vector = embedding_model.encode(text_value, normalize_embeddings=True).tolist()
                new_job.embedding = encode_vector_to_text(vector)
                db.session.commit()
        except Exception as e:
            logger.error(f"Failed to compute job embedding: {str(e)}")

        # Clean up session data
        session.pop('extracted_job_data', None)
        if 'job_file_path' in session:
            filepath = session.pop('job_file_path')
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.error(f"Error removing temporary file {filepath}: {str(e)}")

        flash('Job description saved successfully!', 'success')
        return redirect(url_for('job_descriptions'))

    except Exception as e:
        logger.error(f"Error saving reviewed job description: {str(e)}")
        flash('An unexpected error occurred while saving the job description. Please try again.', 'danger')
        return redirect(url_for('review_job_description'))

@app.route('/ranking/<int:job_id>')
@login_required
def ranking(job_id):
    """View rankings for a job description. Users can only view rankings for their own job descriptions."""
    import time
    start_time = time.time()
    logger.info(f"Starting ranking for job_id {job_id}")

    job = JobDescription.query.get_or_404(job_id)

    # Check if the user owns this job description
    if job.fk_user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to view rankings for this job description.', 'danger')
        return redirect(url_for('job_descriptions'))

    # Allow switching between models for the research experiment
    model_type = request.args.get('model', 'ai_enhanced')
    candidates = Candidate.query.all()
    candidate_dicts = [c.to_dict() for c in candidates]
    job_dict = job.to_dict()
    logger.info(f"Fetched {len(candidate_dicts)} candidates for ranking")

    # Always compute both traditional and AI rankings for comparison
    traditional_ranked_candidates = None
    ai_ranked_candidates = None

    if 'rank_candidates' in globals() and rank_candidates is not None:
        traditional_ranked_candidates = rank_candidates(candidate_dicts, job_dict)

    if model_type == 'baseline':
        ranked_candidates = rank_candidates_baseline(candidate_dicts, job_dict)
        flash('Showing results from the baseline TF-IDF model.', 'info')
    elif model_type == 'traditional':
        ranked_candidates = traditional_ranked_candidates if traditional_ranked_candidates else [{'candidate': c, 'total_score': 0.0} for c in candidate_dicts]
        if not traditional_ranked_candidates:
            flash('Traditional ranking is not available on this server instance.', 'warning')
        else:
            flash('Showing results from the traditional semantic model.', 'info')
    else:  # ai_enhanced (default)
        try:
            from ai_candidate_ranker import rank_candidates_with_ai
            ai_ranked_candidates = rank_candidates_with_ai(candidate_dicts, job_dict)
            ranked_candidates = ai_ranked_candidates
            if AI_ANALYSIS_AVAILABLE and ai_analyzer and ai_analyzer.is_available():
                flash('Showing AI-enhanced rankings with GPT analysis.', 'success')
            else:
                flash('Showing traditional rankings (AI analysis not available - configure OpenAI API key for enhanced features).', 'warning')
        except Exception as e:
            logger.error(f"AI-enhanced ranking failed: {e}")
            # Fallback to traditional ranking
            ranked_candidates = traditional_ranked_candidates if traditional_ranked_candidates else [{'candidate': c, 'total_score': 0.0} for c in candidate_dicts]
            flash('AI-enhanced ranking failed, showing traditional results.', 'warning')

    end_time = time.time()
    logger.info(f"Ranking completed in {end_time - start_time:.2f} seconds for job_id {job_id}")
    return render_template('ranking.html', job=job_dict, ranked_candidates=ranked_candidates, model_type=model_type,
                          traditional_ranked_candidates=traditional_ranked_candidates, ai_ranked_candidates=ai_ranked_candidates)

@app.route('/ranking/<int:job_id>/dashboard')
@login_required
def ranking_dashboard(job_id):
    """Comprehensive analytics dashboard for ranking results."""
    import time
    start_time = time.time()
    logger.info(f"Loading ranking dashboard for job_id {job_id}")

    job = JobDescription.query.get_or_404(job_id)

    # Check if the user owns this job description
    if job.fk_user_id != current_user.id and not current_user.is_admin:
        flash('You do not have permission to view this dashboard.', 'danger')
        return redirect(url_for('job_descriptions'))

    candidates = Candidate.query.all()
    candidate_dicts = [c.to_dict() for c in candidates]
    job_dict = job.to_dict()

    # Get both ranking types for comparison
    traditional_ranked = None
    ai_ranked = None

    if 'rank_candidates' in globals() and rank_candidates is not None:
        traditional_ranked = rank_candidates(candidate_dicts, job_dict)

    try:
        from ai_candidate_ranker import rank_candidates_with_ai
        ai_ranked = rank_candidates_with_ai(candidate_dicts, job_dict)
    except Exception as e:
        logger.error(f"AI ranking failed for dashboard: {e}")

    # Prepare analytics data
    analytics = prepare_ranking_analytics(traditional_ranked, ai_ranked, job_dict)
    logger.info(f"Analytics data prepared: {analytics}")

    end_time = time.time()
    logger.info(f"Ranking dashboard loaded in {end_time - start_time:.2f} seconds for job_id {job_id}")

    return render_template('ranking_dashboard.html', job=job_dict, analytics=analytics)

def prepare_ranking_analytics(traditional_ranked, ai_ranked, job_dict):
    """Prepare comprehensive analytics data for the ranking dashboard."""

    analytics = {
        'score_distribution': {'labels': [], 'ai_scores': [], 'traditional_scores': []},
        'skills_analysis': {'matched': 0, 'missing': 0, 'partial': 0},
        'experience_distribution': {'ranges': ['0-2', '2-5', '5-8', '8+'], 'counts': [0, 0, 0, 0]},
        'education_distribution': {'labels': [], 'counts': []},
        'top_candidates': {'names': [], 'scores': [], 'model': []},
        'correlation_data': {'experience': [], 'scores': []},
        'model_comparison': {'traditional_avg': 0, 'ai_avg': 0, 'difference': 0}
    }

    # Use AI rankings if available, otherwise fall back to traditional
    primary_rankings = ai_ranked if ai_ranked else traditional_ranked

    if primary_rankings:
        # Score distribution for top 10 candidates
        for i, candidate in enumerate(primary_rankings[:10]):
            name = candidate['candidate']['name'][:15] + '...' if len(candidate['candidate']['name']) > 15 else candidate['candidate']['name']
            analytics['score_distribution']['labels'].append(name)

            if ai_ranked:
                analytics['score_distribution']['ai_scores'].append(round(candidate['total_score'] * 100, 1))
            if traditional_ranked and i < len(traditional_ranked):
                analytics['score_distribution']['traditional_scores'].append(round(traditional_ranked[i]['total_score'] * 100, 1))

        # Skills analysis
        total_matched = 0
        total_missing = 0
        for candidate in primary_rankings:
            matched = len(candidate.get('matched_skills', {}).get('matched', []))
            missing = len(candidate.get('matched_skills', {}).get('missing', []))
            total_matched += matched
            total_missing += missing

        analytics['skills_analysis']['matched'] = total_matched
        analytics['skills_analysis']['missing'] = total_missing

        # Experience distribution
        for candidate in primary_rankings:
            exp = candidate['candidate'].get('experience_years', 0)
            if exp < 2:
                analytics['experience_distribution']['counts'][0] += 1
            elif exp < 5:
                analytics['experience_distribution']['counts'][1] += 1
            elif exp < 8:
                analytics['experience_distribution']['counts'][2] += 1
            else:
                analytics['experience_distribution']['counts'][3] += 1

        # Education distribution
        education_counts = {}
        for candidate in primary_rankings:
            education = candidate['candidate'].get('education', [])
            if education:
                degree = education[0].get('degree', 'Unknown')
                education_counts[degree] = education_counts.get(degree, 0) + 1
            else:
                education_counts['No Education Listed'] = education_counts.get('No Education Listed', 0) + 1

        analytics['education_distribution']['labels'] = list(education_counts.keys())
        analytics['education_distribution']['counts'] = list(education_counts.values())

        # Top candidates
        for candidate in primary_rankings[:5]:
            analytics['top_candidates']['names'].append(candidate['candidate']['name'])
            analytics['top_candidates']['scores'].append(round(candidate['total_score'] * 100, 1))
            analytics['top_candidates']['model'].append('AI' if ai_ranked else 'Traditional')

        # Correlation data
        for candidate in primary_rankings:
            analytics['correlation_data']['experience'].append(candidate['candidate'].get('experience_years', 0))
            analytics['correlation_data']['scores'].append(round(candidate['total_score'] * 100, 1))

        # Model comparison
        if ai_ranked:
            ai_avg = sum(c['total_score'] for c in ai_ranked) / len(ai_ranked)
            analytics['model_comparison']['ai_avg'] = round(ai_avg * 100, 1)

        if traditional_ranked:
            trad_avg = sum(c['total_score'] for c in traditional_ranked) / len(traditional_ranked)
            analytics['model_comparison']['traditional_avg'] = round(trad_avg * 100, 1)
            if ai_ranked:
                analytics['model_comparison']['difference'] = round((ai_avg - trad_avg) * 100, 1)

    # Ensure we have at least some default data for charts to render
    if not analytics['score_distribution']['labels']:
        analytics['score_distribution']['labels'] = ['No Data']
        analytics['score_distribution']['ai_scores'] = [0]
        analytics['score_distribution']['traditional_scores'] = [0]

    if not analytics['education_distribution']['labels']:
        analytics['education_distribution']['labels'] = ['No Data']
        analytics['education_distribution']['counts'] = [1]

    return analytics

@limiter.limit("2 per minute")
@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        consent = request.form.get('consent')

        # Validate inputs
        errors = []
        if not consent:
            errors.append('You must agree to the privacy policy and terms of service')
        if not username:
            errors.append('Username is required')
        elif len(username) < 3:
            errors.append('Username must be at least 3 characters long')
        elif User.query.filter_by(username=username).first():
            errors.append('Username already exists')
            
        if not email:
            errors.append('Email is required')
        else:
            try:
                validate_email(email)
                if User.query.filter_by(email=email).first():
                    errors.append('Email already registered')
            except EmailNotValidError as e:
                errors.append('Invalid email format')
                
        if not password:
            errors.append('Password is required')
        elif len(password) < 8:
            errors.append('Password must be at least 8 characters long')
        elif password != confirm_password:
            errors.append('Passwords do not match')
        if errors:
            for error in errors:
                flash(error, 'danger')
            return redirect(url_for('register'))
        try:
            new_user = User(username=username, email=email, is_admin=False)
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            logger.error(f"Registration error for {username}: {str(e)}")
            flash('An error occurred during registration. Please try again.', 'danger')
    return render_template('register.html')

@limiter.limit("3 per minute")
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        remember = request.form.get('remember', False)
        
        # Validate inputs
        if not username or not password:
            flash('Username and password are required', 'danger')
            return render_template('login.html')
            
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user, remember=remember)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        flash('Invalid username or password', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    try:
        username = current_user.username
        logout_user()
        flash(f'You have been successfully logged out, {username}.', 'success')
        logger.info(f"User {username} logged out successfully")
    except Exception as e:
        logger.error(f"Logout error: {str(e)}")
        flash('An error occurred during logout.', 'danger')
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)